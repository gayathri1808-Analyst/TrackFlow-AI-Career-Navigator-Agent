import os
import pypdf
import docx

def read_resume_file(file_path: str) -> str:
    """
    Reads a resume file from the given file path and returns its text content.
    Supports PDF, DOCX, and TXT formats.
    
    Args:
        file_path: The absolute or relative path to the resume file.
        
    Returns:
        The extracted raw text content of the resume, or an error message if extraction fails.
    """
    # Clean file path input
    file_path = file_path.strip().strip("'\"")
    
    if not os.path.exists(file_path):
        return f"Error: File not found at path: {file_path}"
        
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
                
        elif ext == ".pdf":
            text_parts = []
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            extracted = "\n".join(text_parts).strip()
            if not extracted:
                return "Error: PDF file contains no readable text."
            return extracted
            
        elif ext == ".docx":
            doc = docx.Document(file_path)
            text_parts = [para.text for para in doc.paragraphs if para.text]
            # Also extract from tables in docx
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)
            extracted = "\n".join(text_parts).strip()
            if not extracted:
                return "Error: DOCX file contains no readable text."
            return extracted
            
        else:
            return f"Error: Unsupported file format '{ext}'. TrackFlow AI only supports PDF, DOCX, and TXT files."
            
    except Exception as e:
        return f"Error reading file: {e}"
